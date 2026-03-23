"""
Generate system design document for Fire Detection project.
Run: python scripts/generate_doc.py
Output: docs/Fire_Detection_System_Design.docx
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docs"
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "Fire_Detection_System_Design.docx"

# ── Color palette ──────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0xC0, 0x39, 0x2B)   # deep red
C_H1      = RGBColor(0x1A, 0x5F, 0x7A)   # dark teal
C_H2      = RGBColor(0x2E, 0x86, 0xAB)   # medium blue
C_H3      = RGBColor(0x16, 0xA0, 0x85)   # teal-green
C_CODE_BG = RGBColor(0xF4, 0xF4, 0xF4)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER_BG = RGBColor(0x1A, 0x5F, 0x7A)

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def add_page_break(doc: Document):
    doc.add_page_break()

def heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = C_H1
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = C_H2
        run.font.size = Pt(14)
        run.bold = True
    elif level == 3:
        run.font.color.rgb = C_H3
        run.font.size = Pt(12)
        run.bold = True
    return p

def body(doc: Document, text: str, bold_parts: list[str] | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                if idx > 0:
                    p.add_run(remaining[:idx])
                p.add_run(bp).bold = True
                remaining = remaining[idx + len(bp):]
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    return p

def bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p

def code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p

def table_styled(doc: Document, headers: list[str], rows: list[list[str]]):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_HEADER_BG)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = RGBColor(0xEA, 0xF2, 0xF8) if r_idx % 2 == 0 else C_WHITE
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)

    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── COVER PAGE ─────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(72)
run = cover.add_run("HỆ THỐNG PHÁT HIỆN LỬA VÀ KHÓI\nDỰA TRÊN YOLOv10")
run.font.color.rgb = C_TITLE
run.font.size = Pt(26)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("TÀI LIỆU THIẾT KẾ HỆ THỐNG & PIPELINE XÂY DỰNG")
r.font.size = Pt(13)
r.font.color.rgb = C_H2
r.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    f"Phiên bản: 1.0  |  Ngày: {datetime.date.today().strftime('%d/%m/%Y')}\n"
    "Dataset: D-Fire (21,527 ảnh)  |  Model: YOLOv10x\n"
    "Framework: Ultralytics · MLflow · FastAPI · Docker"
).font.size = Pt(11)

add_page_break(doc)

# ── 1. TỔNG QUAN DỰ ÁN ────────────────────────────────────────────────────
heading(doc, "1. TỔNG QUAN DỰ ÁN", 1)
body(doc,
    "Dự án xây dựng hệ thống phát hiện lửa và khói thời gian thực sử dụng mô hình "
    "YOLOv10x — một trong những kiến trúc object detection tiên tiến nhất. "
    "Hệ thống được thiết kế theo hướng production-ready với đầy đủ MLOps pipeline: "
    "từ quản lý dataset, fine-tuning model, theo dõi thực nghiệm (MLflow), đến "
    "triển khai API inference với khả năng stream camera RTSP trực tiếp.",
    bold_parts=["YOLOv10x", "MLflow", "RTSP"])

heading(doc, "1.1 Mục tiêu", 2)
bullet(doc, "Phát hiện lửa (fire) và khói (smoke) trong video stream thời gian thực")
bullet(doc, "Fine-tune YOLOv10x trên bộ D-Fire với 21,527 ảnh đã được label")
bullet(doc, "Xây dựng MLOps pipeline hoàn chỉnh: tracking thực nghiệm, artifact versioning")
bullet(doc, "API inference hỗ trợ đa camera RTSP/webcam qua WebSocket")
bullet(doc, "Đạt mAP50 > 0.75 sau fine-tuning (baseline: 0.601)")

heading(doc, "1.2 Thông số kỹ thuật chính", 2)
table_styled(doc,
    ["Thành phần", "Chi tiết"],
    [
        ["Model backbone",      "YOLOv10x (Extra-large)"],
        ["Init weights",        "Model/best.pt — HuggingFace TommyNgx/YOLOv10-Fire-and-Smoke"],
        ["Baseline mAP50",      "0.601 (trước fine-tuning)"],
        ["Dataset",             "D-Fire: 21,527 ảnh (train=14,122 / val=3,099 / test=4,306)"],
        ["Classes",             "fire (id=0), smoke (id=1)"],
        ["Training strategy",   "Progressive freeze fine-tuning (freeze=10 backbone layers)"],
        ["Tracking",            "MLflow + PostgreSQL + MinIO (Docker Compose)"],
        ["Inference API",       "FastAPI + WebSocket + OpenCV RTSP"],
        ["Python version",      "3.10+"],
        ["GPU yêu cầu",        "8GB+ VRAM (RTX 3070/4070 trở lên)"],
    ]
)

# ── 2. KIẾN TRÚC HỆ THỐNG ─────────────────────────────────────────────────
heading(doc, "2. KIẾN TRÚC HỆ THỐNG", 1)
body(doc,
    "Hệ thống được tổ chức thành 7 layer độc lập, giao tiếp với nhau qua "
    "dependency injection và config-driven design. Không có hardcoded path "
    "hay credential nào trong source code.",
    bold_parts=["7 layer", "config-driven design"])

heading(doc, "2.1 Sơ đồ tổng thể", 2)
code_block(doc,
"""┌─────────────────────────────────────────────────────────────────────┐
│                        FIRE DETECTION SYSTEM                        │
├──────────────┬──────────────┬──────────────┬────────────────────────┤
│   config/    │    src/      │   infra/     │      data/             │
│  (7 YAML)    │  (8 module)  │  (Docker)    │   (D-Fire 21K img)     │
├──────────────┴──────────────┴──────────────┴────────────────────────┤
│                                                                     │
│  ┌──────────────────────────────────────────────────────────────┐  │
│  │                    PIPELINE ORCHESTRATOR                     │  │
│  │   Stage1 → Stage2 → Stage3 → ... → Stage8 → Stage9(Train)  │  │
│  └──────────────────────────────────────────────────────────────┘  │
│                              │                                      │
│              ┌───────────────┼───────────────┐                     │
│              ▼               ▼               ▼                     │
│      ┌──────────────┐ ┌──────────┐ ┌────────────────┐             │
│      │ MLflow       │ │  YOLO    │ │   FastAPI      │             │
│      │ Tracking     │ │ Training │ │   Inference    │             │
│      │ (Postgres +  │ │ (Ultra-  │ │   API +        │             │
│      │  MinIO)      │ │  lytics) │ │   WebSocket    │             │
│      └──────────────┘ └──────────┘ └────────────────┘             │
└─────────────────────────────────────────────────────────────────────┘""")

heading(doc, "2.2 Cấu trúc thư mục", 2)
code_block(doc,
"""D:\\MSA_28\\Final_project_Fire_Detection\\
├── config/                   # 7 YAML files — toàn bộ config nằm đây
│   ├── app.yaml              # Metadata: tên project, env, logging
│   ├── api.yaml              # FastAPI: host, port, thresholds, auth
│   ├── dataset.yaml          # Dataset: source, path, data_yaml_name
│   ├── environment.yaml      # Python version constraints
│   ├── mlflow.yaml           # MLflow URI, MinIO endpoint/bucket
│   ├── model.yaml            # Init weights: source, repo, local path
│   └── training.yaml         # Hyperparams: epochs, lr, freeze, resume
│
├── src/                      # Python source (28 files, 8 packages)
│   ├── core/                 # Tiện ích lõi
│   ├── dataset/              # Xử lý dataset
│   ├── model/                # Load model
│   ├── training/             # Training wrapper
│   ├── tracking/             # MLflow integration
│   ├── pipeline/             # Orchestrator + stages
│   ├── reporting/            # Báo cáo thực nghiệm
│   └── api/                  # FastAPI inference server
│
├── infra/mlflow/             # Docker Compose stack
│   ├── docker-compose.yml    # PostgreSQL + MinIO + MLflow
│   └── Dockerfile
│
├── data/DFire/               # D-Fire dataset (3.0 GB)
├── Model/best.pt             # Init weights YOLOv10x (62 MB)
├── runs/train/               # Training outputs (Ultralytics format)
├── reports/                  # Generated reports (JSON + MD)
├── logs/                     # App logs (rotating, 10MB × 3)
├── scripts/                  # Utility scripts
├── .env                      # Secrets (không commit vào git)
└── .env.example              # Template cho team""")

heading(doc, "2.3 Nguyên tắc thiết kế", 2)
table_styled(doc,
    ["Nguyên tắc", "Áp dụng trong dự án"],
    [
        ["Config-driven",       "Toàn bộ path/URI/credential đọc từ config/*.yaml và .env — zero hardcode"],
        ["Separation of Concerns", "8 package riêng biệt, mỗi cái một trách nhiệm duy nhất"],
        ["Dependency Injection","ConfigLoader được inject vào tất cả class thay vì singleton global"],
        ["Graceful Degradation","Pipeline tiếp tục dù MLflow/dataset unavailable — non-fatal stages"],
        ["Reproducibility",    "Env fingerprint + git commit + hyperparams đều log vào MLflow"],
        ["Security",           "Credentials từ .env, RTSP URL validation, Bearer token auth API"],
    ]
)

add_page_break(doc)

# ── 3. PIPELINE XÂY DỰNG ──────────────────────────────────────────────────
heading(doc, "3. PIPELINE XÂY DỰNG (BUILD PIPELINE)", 1)
body(doc,
    "Pipeline được khởi chạy bằng lệnh: python -m src.pipeline.orchestrator\n"
    "Gồm 9 stage chạy tuần tự, chia sẻ context dict. Mỗi stage trả về True (thành công) "
    "hoặc False (cảnh báo không nghiêm trọng).",
    bold_parts=["python -m src.pipeline.orchestrator", "9 stage", "True", "False"])

heading(doc, "3.1 Tổng quan các Stage", 2)
table_styled(doc,
    ["#", "Stage", "Class", "Output", "Fatal?"],
    [
        ["1", "LoadConfig",           "LoadConfigStage",           "Config dict trong context",                "Có"],
        ["2", "EnvFingerprint",       "EnvFingerprintStage",       "env_fingerprint.json",                    "Không"],
        ["3", "InitWeightsManifest",  "InitWeightsManifestStage",  "init_weights_manifest.json + SHA256",     "Không"],
        ["4", "DatasetValidate",      "DatasetValidateStage",      "ValidationReport",                        "Không"],
        ["5", "DatasetAnalyze",       "DatasetAnalyzeStage",       "DatasetStats (có cache)",                 "Không"],
        ["6", "DatasetReportExport",  "DatasetReportExportStage",  "dataset_report.md + dataset_stats.json",  "Không"],
        ["7", "MLflowSmokeRun",       "MLflowSmokeRunStage",       "MLflow run + framework artifacts",        "Không"],
        ["8", "ExperimentSummary",    "ExperimentSummaryExportStage","experiment_summary.md",                 "Không"],
        ["9", "Train",                "TrainStage",                "Trained model + per-epoch MLflow metrics","Không"],
    ]
)

heading(doc, "3.2 Luồng xử lý chi tiết", 2)
code_block(doc,
"""orchestrator.py
     │
     ├─► Stage 1: LoadConfigStage
     │     └─ Đọc 7 YAML files, validate required keys
     │     └─ Nạp .env qua python-dotenv (credentials)
     │     └─ ctx ← {app_cfg, dataset_cfg, mlflow_cfg, model_cfg, env_cfg}
     │
     ├─► Stage 2: EnvFingerprintStage
     │     └─ Thu thập: OS, Python, GPU (CUDA), pip packages, env vars
     │     └─ Gắn git commit/branch/dirty flag
     │     └─ Xuất: reports/evaluation/env_fingerprint.json
     │
     ├─► Stage 3: InitWeightsManifestStage
     │     └─ Kiểm tra Model/best.pt tồn tại
     │     └─ Tính SHA256 của file weights
     │     └─ Xuất: reports/evaluation/init_weights_manifest.json
     │
     ├─► Stage 4: DatasetValidateStage
     │     └─ Load data.yaml từ data/DFire/
     │     └─ Validate: images/ + labels/ tồn tại, 1:1 correspondence
     │     └─ Kiểm tra ảnh không bị corrupt (PIL.Image.verify)
     │     └─ Non-fatal nếu dataset không có
     │
     ├─► Stage 5: DatasetAnalyzeStage
     │     └─ Cache check: dataset_stats.json mtime > data.yaml mtime?
     │     │     └─ HIT: load từ JSON (~1ms, bỏ qua label scan)
     │     │     └─ MISS: đọc 21,527 label files, tính class distribution
     │     └─ DatasetStats: images/split, boxes/image, class dist
     │
     ├─► Stage 6: DatasetReportExportStage
     │     └─ Xuất: reports/evaluation/dataset_report.md (Markdown)
     │     └─ Xuất: reports/evaluation/dataset_stats.json (JSON cache)
     │
     ├─► Stage 7: MLflowSmokeRunStage
     │     └─ Kết nối MLflow server (localhost:5000)
     │     └─ Tạo run "framework-smoke", log params + metrics
     │     └─ Upload 4 framework artifacts lên MinIO
     │     └─ Non-fatal nếu MLflow không chạy
     │
     ├─► Stage 8: ExperimentSummaryExportStage
     │     └─ Query 20 runs gần nhất từ MLflow
     │     └─ Xuất: reports/experiments/experiment_summary.md
     │
     └─► Stage 9: TrainStage (enabled: true/false trong training.yaml)
           └─ Resume check: resume=true → tìm last.pt mới nhất
           └─ Load model từ best.pt (hoặc checkpoint)
           └─ Tạo temp data.yaml với absolute path
           └─ Disable Ultralytics MLflow callback (tránh duplicate runs)
           └─ Đăng ký on_fit_epoch_end callback → log ALL metrics mỗi epoch
           └─ YOLO.train() với toàn bộ hyperparams từ training.yaml
           └─ Log best.pt + training plots → MLflow artifacts
           └─ Log final mAP50, mAP50-95, precision, recall""")

heading(doc, "3.3 Pattern BaseStage", 2)
body(doc,
    "Tất cả stages kế thừa BaseStage(ABC) — đảm bảo interface nhất quán:",
    bold_parts=["BaseStage(ABC)"])
code_block(doc,
"""class BaseStage(ABC):
    name: str = "base"

    def __init__(self, cfg: ConfigLoader) -> None:
        self.cfg = cfg

    @abstractmethod
    def run(self, ctx: dict[str, Any]) -> bool:
        ...   # True = success, False = non-fatal warning""")

add_page_break(doc)

# ── 4. MODULE CHI TIẾT ────────────────────────────────────────────────────
heading(doc, "4. MÔ TẢ CHI TIẾT CÁC MODULE", 1)

heading(doc, "4.1 src/core — Tiện ích lõi", 2)
table_styled(doc,
    ["File", "Class/Function", "Mô tả"],
    [
        ["config_loader.py", "ConfigLoader",        "Load YAML configs, validate required keys, load .env qua dotenv, cache kết quả"],
        ["logger.py",        "get_logger(name)",    "Logging đôi: console + rotating file (10MB × 3 backup). Format: timestamp|level|module|msg"],
        ["env_fingerprint.py","collect(), export()","Thu thập OS/Python/GPU/pip/env vars. Xuất JSON cho reproducibility"],
        ["git_info.py",      "get_info()",          "Lấy git commit hash, branch, dirty flag. Fallback nếu không phải git repo"],
    ]
)

heading(doc, "4.2 src/dataset — Xử lý Dataset", 2)
table_styled(doc,
    ["File", "Class/Function", "Mô tả"],
    [
        ["metadata.py",       "load_metadata()",    "Parse data.yaml: class names, num_classes, split paths (absolute)"],
        ["validator.py",      "validate()",         "Kiểm tra YOLOv8 layout: images/ + labels/, 1:1 pairing, PIL verify"],
        ["analyzer.py",       "analyze()",          "Tính stats: images/split, boxes/image (mean/med/min/max/std), class dist. Có cache JSON"],
        ["report_exporter.py","export()",           "Xuất dataset_report.md (Markdown) + dataset_stats.json (JSON)"],
    ]
)

heading(doc, "4.3 src/training — Training Wrapper", 2)
body(doc,
    "FireDetectionTrainer bọc YOLO.train() với đầy đủ MLflow integration:",
    bold_parts=["FireDetectionTrainer"])
bullet(doc, "Resume support: tự tìm last.pt mới nhất hoặc dùng path chỉ định")
bullet(doc, "Tắt Ultralytics MLflow callback gốc để tránh duplicate run")
bullet(doc, "Callback on_fit_epoch_end: log TẤT CẢ metrics động từ trainer.metrics mỗi epoch")
bullet(doc, "Tạo temp data.yaml với absolute path (tránh Ultralytics path resolution bug)")
bullet(doc, "Log hyperparams đầy đủ: epochs, batch, imgsz, freeze, lr0, lrf, warmup, augmentation")
bullet(doc, "Log artifacts: best.pt, results.png, confusion_matrix.png, PR_curve.png, F1_curve.png")

heading(doc, "4.4 src/tracking — MLflow Integration", 2)
table_styled(doc,
    ["File", "Mô tả"],
    [
        ["mlflow_client.py",      "Wrapper MLflow SDK. Đọc credentials từ env (MINIO_ACCESS_KEY/SECRET_KEY). Context manager run lifecycle"],
        ["artifact_manager.py",   "Batch upload framework artifacts (4 files JSON/MD) lên MinIO bucket"],
        ["experiment_manager.py", "Query runs: get_last_n_runs(), get_best_run(metric). Dùng cho summary report"],
    ]
)

heading(doc, "4.5 src/api — Inference API", 2)
table_styled(doc,
    ["File", "Mô tả"],
    [
        ["main.py",         "FastAPI app: REST + WebSocket endpoints. Bearer token auth (optional). RTSP URL validation. Lifespan management"],
        ["detector.py",     "Thread-safe YOLO inference. Input: BGR frame. Output: list[{class, confidence, bbox}]"],
        ["stream_manager.py","Quản lý đa camera. FrameGrabber loại bỏ FFmpeg buffer lag. Exponential backoff reconnect. Per-subscriber queue"],
    ]
)

add_page_break(doc)

# ── 5. CẤU HÌNH HỆ THỐNG ─────────────────────────────────────────────────
heading(doc, "5. CẤU HÌNH HỆ THỐNG", 1)

heading(doc, "5.1 config/training.yaml — Hyperparameters", 2)
code_block(doc,
"""enabled: true           # false để skip TrainStage
run_name: dfire_finetune

# Resume interrupted training
resume: false           # true = tiếp tục từ checkpoint
resume_from: null       # null = auto-detect last.pt mới nhất

# Epochs & size
epochs: 50              # 5 để test pipeline, 50+ để train thật
imgsz: 640
batch: 4                # YOLOv10x cần 8-12GB VRAM

# Device
device: "0"             # GPU index, hoặc "cpu"
workers: 4

# Fine-tuning strategy (progressive freeze)
freeze: 10              # Đóng băng 10 layers đầu (backbone)
lr0: 0.0005             # LR thấp — fine-tuning, không train từ đầu
lrf: 0.01               # Final LR = lr0 × lrf = 0.000005
momentum: 0.937
weight_decay: 0.0005
warmup_epochs: 2.0

# Augmentation (nhẹ hơn train từ đầu)
hsv_h: 0.015  hsv_s: 0.7  hsv_v: 0.4
flipud: 0.0   fliplr: 0.5
mosaic: 1.0   mixup: 0.0

patience: 20            # Early stopping""")

heading(doc, "5.2 config/api.yaml — Inference API", 2)
code_block(doc,
"""host: "0.0.0.0"
port: 8000
model_path: "runs/train/.../weights/best.pt"
model_fallback_path: "Model/best.pt"
confidence_threshold: 0.25
iou_threshold: 0.45
target_fps: 15
jpeg_quality: 75
frame_width: 640
frame_height: 480
alert_cooldown_sec: 3
api_key: null           # null = auth disabled (dev), set via API_KEY env var""")

heading(doc, "5.3 .env — Secrets (không commit vào git)", 2)
code_block(doc,
"""# MinIO credentials (MLflow artifact store)
MINIO_ACCESS_KEY=minioadmin
MINIO_SECRET_KEY=minioadmin

# API bearer token (write endpoints)
API_KEY=                # để trống = auth disabled""")

add_page_break(doc)

# ── 6. MLFLOW TRACKING ────────────────────────────────────────────────────
heading(doc, "6. MLFLOW TRACKING INFRASTRUCTURE", 1)

heading(doc, "6.1 Kiến trúc MLflow Stack", 2)
code_block(doc,
"""Docker Compose (infra/mlflow/docker-compose.yml):

┌─────────────────────────────────────────────────────┐
│  MLflow Server (:5000)                              │
│    └─ Backend Store: PostgreSQL (:5432)             │
│    └─ Artifact Store: MinIO (:9000)                 │
├─────────────────────────────────────────────────────┤
│  MinIO Console: http://localhost:9001               │
│    Bucket: mlflow                                   │
│    Credentials: từ .env (MINIO_ACCESS_KEY/SECRET)   │
└─────────────────────────────────────────────────────┘

Khởi động:
  cd infra/mlflow && docker compose up -d""")

heading(doc, "6.2 Thông tin được log vào MLflow", 2)
table_styled(doc,
    ["Loại", "Nội dung"],
    [
        ["Params",   "epochs, batch, imgsz, freeze, lr0, lrf, warmup, patience, device, dataset, weights, resumed"],
        ["Metrics",  "Mỗi epoch: train.box_loss, train.cls_loss, train.dfl_loss, val.*, metrics.mAP50, metrics.mAP50-95, precision, recall"],
        ["Tags",     "phase=train, stage=finetune, dataset=DFire, freeze_layers=10"],
        ["Artifacts","best.pt, last.pt, results.png, confusion_matrix.png, PR_curve.png, F1_curve.png, dataset_report.md, env_fingerprint.json"],
    ]
)

heading(doc, "6.3 MLflow Client Usage", 2)
code_block(doc,
"""from src.tracking.mlflow_client import MLflowClient

client = MLflowClient(config_loader=cfg)
with client.run(run_name="dfire_finetune_20260305", tags={"phase": "train"}):
    client.log_params({"lr0": 0.0005, "batch": 4, "freeze": 10})
    client.log_metrics({"mAP50": 0.754}, step=epoch)
    client.log_artifact("runs/train/.../weights/best.pt", "weights")""")

add_page_break(doc)

# ── 7. INFERENCE API ──────────────────────────────────────────────────────
heading(doc, "7. INFERENCE API", 1)

heading(doc, "7.1 REST Endpoints", 2)
table_styled(doc,
    ["Method", "Endpoint", "Auth?", "Mô tả"],
    [
        ["GET",    "/",                   "Không", "Dashboard UI (HTML)"],
        ["GET",    "/api/config",         "Không", "Lấy detection thresholds"],
        ["PUT",    "/api/config",         "Có",    "Cập nhật confidence threshold"],
        ["GET",    "/api/cameras",        "Không", "Danh sách camera đang stream"],
        ["GET",    "/api/webcams",        "Không", "Scan webcam devices (0-9)"],
        ["POST",   "/api/cameras",        "Có",    "Thêm camera RTSP hoặc webcam"],
        ["DELETE", "/api/cameras/{id}",   "Có",    "Xoá camera stream"],
        ["WS",     "/ws/{camera_id}",     "Không", "Live stream + detections JSON"],
    ]
)

heading(doc, "7.2 RTSP Stream Architecture", 2)
body(doc,
    "Kiến trúc streaming được thiết kế để đạt latency thấp nhất có thể:",
    bold_parts=["latency thấp"])
code_block(doc,
"""Camera (RTSP) ──► FrameGrabber Thread ──► Latest Frame Buffer
                        │                         │
                        │  drain full speed        │
                        │  keep only latest        │
                        ▼                         ▼
                Processing Thread ◄──────── get fresh frame
                        │
                        ├─► YOLO detect (~30ms)
                        ├─► Annotate bounding boxes
                        ├─► JPEG encode (quality=75)
                        └─► Broadcast → asyncio.Queue(maxsize=4) per subscriber
                                              │
                                              ▼
                                    WebSocket Client(s)""")

bullet(doc, "FrameGrabber loại bỏ FFmpeg buffer lag: latency giảm từ 500ms–2s xuống còn <50ms")
bullet(doc, "Mỗi subscriber có queue riêng, maxsize=4 — drop oldest nếu subscriber chậm")
bullet(doc, "OpenCV CAP_PROP_BUFFERSIZE=1 — chỉ giữ 1 frame trong OS buffer")
bullet(doc, "Exponential backoff reconnect: delay 2s × 1.5x, max 30s, unlimited attempts")
bullet(doc, "FFmpeg options: nobuffer, low_delay, framedrop, stimeout=5s, tcp transport")

heading(doc, "7.3 Khởi động API", 2)
code_block(doc,
"""# Cách 1 — Module trực tiếp
python -m src.api.main

# Cách 2 — Uvicorn với auto-reload (development)
uvicorn src.api.main:app --host 0.0.0.0 --port 8000 --reload

# Gọi API với authentication
curl -X POST http://localhost:8000/api/cameras \\
  -H "Authorization: Bearer <API_KEY>" \\
  -H "Content-Type: application/json" \\
  -d '{"type":"rtsp","url":"rtsp://192.168.1.100:554/stream","name":"Camera 1"}'""")

add_page_break(doc)

# ── 8. DATASET ────────────────────────────────────────────────────────────
heading(doc, "8. DATASET — D-FIRE", 1)

heading(doc, "8.1 Thông số Dataset", 2)
table_styled(doc,
    ["Thuộc tính", "Giá trị"],
    [
        ["Tên",         "D-Fire (Diverse Fire Dataset)"],
        ["Tổng ảnh",    "21,527 ảnh"],
        ["Train split", "14,122 ảnh (65.6%)"],
        ["Val split",   "3,099 ảnh (14.4%)"],
        ["Test split",  "4,306 ảnh (20.0%)"],
        ["Classes",     "2: fire (id=0), smoke (id=1)"],
        ["Format",      "YOLOv8 (normalized xywh)"],
        ["Location",    "data/DFire/data/{train,val,test}/{images,labels}/"],
        ["Class remap", "Original: smoke=0, fire=1 → Remapped: fire=0, smoke=1 (khớp best.pt)"],
    ]
)

heading(doc, "8.2 Cấu trúc thư mục Dataset", 2)
code_block(doc,
"""data/DFire/
├── data.yaml                    # YOLO metadata (path, nc, names)
└── data/
    ├── train/
    │   ├── images/              (14,122 files: .jpg/.png)
    │   └── labels/              (14,122 files: .txt, YOLO format)
    ├── val/
    │   ├── images/              (3,099 files)
    │   └── labels/              (3,099 files)
    └── test/
        ├── images/              (4,306 files)
        └── labels/              (4,306 files)

Mỗi label file (.txt): mỗi dòng = 1 object
Format: <class_id> <x_center> <y_center> <width> <height>  (normalized 0-1)
Ví dụ:  0 0.512 0.341 0.234 0.189  ← fire object""")

heading(doc, "8.3 Dataset Caching", 2)
body(doc,
    "Dataset analysis (đọc 21,527 label files) được cache tự động vào "
    "reports/evaluation/dataset_stats.json. Cache valid khi mtime của JSON "
    "mới hơn mtime của data.yaml. Cache hit giúp orchestrator bỏ qua label scan "
    "hoàn toàn (~30-60s → <1ms).",
    bold_parts=["dataset_stats.json", "~30-60s → <1ms"])

add_page_break(doc)

# ── 9. TRAINING STRATEGY ──────────────────────────────────────────────────
heading(doc, "9. CHIẾN LƯỢC TRAINING", 1)

heading(doc, "9.1 Progressive Freeze Fine-Tuning", 2)
body(doc,
    "Thay vì train từ đầu (from scratch), dự án dùng progressive freeze fine-tuning "
    "để tránh catastrophic forgetting và tận dụng features đã học của YOLOv10x:",
    bold_parts=["progressive freeze fine-tuning", "catastrophic forgetting"])

table_styled(doc,
    ["Giai đoạn", "freeze", "lr0", "Epochs", "Mục tiêu"],
    [
        ["Stage 1 (hiện tại)", "10 layers (backbone)", "0.0005", "50",   "Chỉ train detection head"],
        ["Stage 2 (kế hoạch)", "5 layers (cuối backbone)", "0.0001", "30", "Unfreeze neck + cuối backbone"],
        ["Stage 3 (kế hoạch)", "0 (full unfreeze)",    "0.00005", "20",  "Fine-tune toàn bộ mô hình"],
    ]
)

heading(doc, "9.2 Resume Training", 2)
body(doc,
    "Nếu training bị interrupt (mất điện, OOM), có thể tiếp tục từ checkpoint:",
    bold_parts=["tiếp tục từ checkpoint"])
code_block(doc,
"""# Trong config/training.yaml:
resume: true
resume_from: null    # null = tự tìm last.pt mới nhất trong runs/train/

# Hoặc chỉ định explicit:
resume_from: "runs/train/dfire_finetune_20260305_120000/weights/last.pt"

# Logic trong trainer.py:
# 1. resume=true → tìm last.pt mới nhất (sort by mtime)
# 2. Tìm thấy → load last.pt thay vì best.pt
# 3. YOLO.train(resume=True) → restore optimizer state + epoch count
# 4. Không tìm thấy → warning + fallback về init weights""")

heading(doc, "9.3 Ước tính thời gian training", 2)
table_styled(doc,
    ["Cấu hình", "Thời gian/epoch", "50 epochs", "VRAM cần"],
    [
        ["YOLOv10x, batch=4, imgsz=640, GPU RTX 3080",  "~3-4 phút", "~2.5-3 giờ", "8GB"],
        ["YOLOv10x, batch=8, imgsz=640, GPU RTX 4090",  "~1.5-2 phút", "~1.5 giờ",   "12GB"],
        ["YOLOv10x, batch=4, imgsz=640, CPU only",       "~45-60 phút","~40 giờ",     "RAM 16GB+"],
    ]
)

add_page_break(doc)

# ── 10. SECURITY ──────────────────────────────────────────────────────────
heading(doc, "10. BẢO MẬT", 1)

heading(doc, "10.1 Các lớp bảo mật đã implement", 2)
table_styled(doc,
    ["Lớp", "Vấn đề", "Giải pháp"],
    [
        ["Credentials", "Hardcoded minioadmin trong source code",          "Load từ .env qua python-dotenv. .env excluded khỏi git"],
        ["API Auth",    "Mọi người đều add/remove camera được",            "Bearer token dependency (_require_auth). Disabled theo mặc định, bật bằng API_KEY env var"],
        ["RTSP URL",    "SSRF: user inject file:// hoặc http:// URL",     "Pydantic field_validator: chỉ rtsp:// hoặc rtsps://, phải có hostname"],
        ["Input type",  "type field không validate",                       "field_validator('type'): chỉ 'rtsp' hoặc 'webcam'"],
        ["HTTP Error",  "return {}, 400 sai cú pháp FastAPI",             "raise HTTPException(status_code=400) đúng chuẩn"],
    ]
)

heading(doc, "10.2 Hướng dẫn bảo mật cho Production", 2)
bullet(doc, "Đặt API_KEY=<random-256bit-hex> trong .env hoặc CI/CD secrets")
bullet(doc, "Thay mật khẩu MinIO (MINIO_ACCESS_KEY/MINIO_SECRET_KEY) trước khi deploy")
bullet(doc, "Thêm HTTPS/TLS phía trước FastAPI (nginx reverse proxy hoặc Cloudflare Tunnel)")
bullet(doc, "Giới hạn CORS origins trong production (thay vì allow all)")
bullet(doc, "Đặt logging_level: WARNING trong config/app.yaml cho production")

add_page_break(doc)

# ── 11. OUTPUTS ───────────────────────────────────────────────────────────
heading(doc, "11. CÁC OUTPUT ĐƯỢC TẠO RA", 1)

table_styled(doc,
    ["File", "Tạo bởi Stage", "Mô tả"],
    [
        ["reports/evaluation/env_fingerprint.json",      "EnvFingerprint",   "OS, Python, GPU, pip packages, env vars"],
        ["reports/evaluation/init_weights_manifest.json","InitWeightsManifest","SHA256, size, source của best.pt"],
        ["reports/evaluation/dataset_report.md",         "DatasetReportExport","Markdown report: classes, splits, box stats"],
        ["reports/evaluation/dataset_stats.json",        "DatasetReportExport","JSON cache: class distribution, images/split"],
        ["reports/experiments/experiment_summary.md",    "ExperimentSummary", "Bảng 20 MLflow runs gần nhất"],
        ["runs/train/<name>/weights/best.pt",            "TrainStage",       "Model tốt nhất theo val mAP50"],
        ["runs/train/<name>/weights/last.pt",            "TrainStage",       "Checkpoint cuối để resume"],
        ["runs/train/<name>/results.csv",                "Ultralytics",      "Per-epoch metrics CSV"],
        ["runs/train/<name>/results.png",                "Ultralytics",      "Training curves plot"],
        ["runs/train/<name>/confusion_matrix.png",       "Ultralytics",      "Ma trận nhầm lẫn"],
        ["logs/app.log",                                 "Logger",           "Application logs (rotating 10MB × 3)"],
    ]
)

add_page_break(doc)

# ── 12. HƯỚNG DẪN VẬN HÀNH ───────────────────────────────────────────────
heading(doc, "12. HƯỚNG DẪN VẬN HÀNH", 1)

heading(doc, "12.1 Khởi động lần đầu", 2)
code_block(doc,
"""# 1. Tạo môi trường ảo và cài dependencies
python -m venv .venv
.venv\\Scripts\\activate      # Windows
pip install -r requirements.txt
pip install ultralytics torch torchvision  # cho training

# 2. Cấu hình secrets
cp .env.example .env
# Sửa .env nếu cần (mặc định minioadmin phù hợp local dev)

# 3. Khởi động MLflow infrastructure
cd infra/mlflow && docker compose up -d && cd ../..

# 4. Tạo bucket mlflow trong MinIO console
# Truy cập http://localhost:9001 (minioadmin/minioadmin)
# Tạo bucket tên: mlflow

# 5. Chạy pipeline (framework test, training disabled)
python -m src.pipeline.orchestrator

# 6. Bật training và chạy lại
# Trong config/training.yaml: enabled: true, epochs: 50
python -m src.pipeline.orchestrator""")

heading(doc, "12.2 Xem kết quả trên MLflow UI", 2)
code_block(doc,
"""# Mở MLflow UI
http://localhost:5000

# Experiment: fire-detection
# Xem metrics: mAP50, mAP50-95, precision, recall theo epoch
# Xem artifacts: best.pt, training plots
# So sánh runs: chọn nhiều runs → Compare""")

heading(doc, "12.3 Chạy Inference API", 2)
code_block(doc,
"""# Đảm bảo model_path trong config/api.yaml trỏ đúng best.pt
python -m src.api.main

# Dashboard: http://localhost:8000
# Thêm camera RTSP (với auth nếu đã set API_KEY):
curl -X POST http://localhost:8000/api/cameras \\
  -H "Authorization: Bearer <API_KEY>" \\
  -d '{"type":"rtsp","url":"rtsp://192.168.1.100:554/stream"}'""")

# ── SAVE ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
